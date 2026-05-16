from __future__ import annotations
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config
from templates.docx_styles import STYLES


# ── Low-level Word XML helpers ────────────────────────────────────────────────

def _add_page_numbers(section) -> None:
    # Inserts a `PAGE` field code in the header so Word auto-numbers pages;
    # different_first_page_header_footer=True means the title page has no visible number
    section.different_first_page_header_footer = True
    hdr = section.header
    p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_end)


def _apply_page_setup(doc: Document) -> None:
    # Applies A4 dimensions and Juridical College margins from STYLES constants
    s = doc.sections[0]
    s.page_width = STYLES["page"]["width"]
    s.page_height = STYLES["page"]["height"]
    s.left_margin = STYLES["page"]["margin_left"]
    s.right_margin = STYLES["page"]["margin_right"]
    s.top_margin = STYLES["page"]["margin_top"]
    s.bottom_margin = STYLES["page"]["margin_bottom"]
    _add_page_numbers(s)


def _set_run_font(run, size: Pt, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic


# ── Paragraph-level builders ──────────────────────────────────────────────────

def _add_body_paragraph(doc: Document, text: str) -> None:
    # Standard body paragraph: justified, 15 mm first-line indent, 1.5 line spacing
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.first_line_indent = Mm(15)
    fmt.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_font(run, Pt(12))


def _add_chapter_heading(doc: Document, text: str) -> None:
    # Chapter headings always start on a new page (page_break_before=True)
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    _set_run_font(run, Pt(14), bold=True)


def _add_subchapter_heading(doc: Document, text: str) -> None:
    # Subchapter headings use small caps via direct XML manipulation
    # (python-docx has no built-in small caps property)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    _set_run_font(run, Pt(12))
    rpr = run._r.get_or_add_rPr()
    small_caps = OxmlElement("w:smallCaps")
    small_caps.set(qn("w:val"), "true")
    rpr.append(small_caps)


# ── Page / section builders ───────────────────────────────────────────────────

def _add_title_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("JURIDICAL COLLEGE")
    _set_run_font(run, Pt(16))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(48)
    run2 = p2.add_run(config.PAPER_TITLE.upper())
    _set_run_font(run2, Pt(28))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    run3 = p3.add_run("QUALIFICATION PAPER")
    _set_run_font(run3, Pt(16))

    for _ in range(4):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run4 = p4.add_run("Author: _________________________ [name, surname, group code]")
    _set_run_font(run4, Pt(16))

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run5 = p5.add_run("Scientific Supervisor: _______________ [degree, name, surname]")
    _set_run_font(run5, Pt(16))

    for _ in range(4):
        doc.add_paragraph()

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run("RIGA 2026")
    _set_run_font(run6, Pt(16))


def _add_toc_placeholder(doc: Document) -> None:
    # Word auto-generates the TOC when the user presses F9 after opening the file
    _add_chapter_heading(doc, "TABLE OF CONTENTS")
    _add_body_paragraph(doc, "[Table of contents — update field in Word after opening]")


def _add_abbreviations(doc: Document, abbreviations: list[str]) -> None:
    # Only included when the paper has 7+ distinct abbreviations (Juridical College rule)
    _add_chapter_heading(doc, "LIST OF ABBREVIATIONS")
    for abbr in abbreviations:
        _add_body_paragraph(doc, f"{abbr} — [expand abbreviation]")
    doc.add_page_break()


def _add_section_content(doc: Document, section_id: str, text: str) -> None:
    # Introduction and conclusions use chapter-level headings; subchapters use small caps
    meta = config.SECTION_META.get(section_id, {})
    heading = meta.get("heading", section_id)
    kind = meta.get("kind", "subchapter")

    if kind in ("introduction", "conclusions"):
        _add_chapter_heading(doc, heading)
    else:
        _add_subchapter_heading(doc, heading)

    for paragraph in text.split("\n\n"):
        stripped = paragraph.strip()
        if stripped:
            _add_body_paragraph(doc, stripped)


def _add_bibliography(doc: Document, bibliography: list[dict]) -> None:
    # Hanging indent (left_indent 10 mm, first_line_indent −10 mm) per spec
    _add_chapter_heading(doc, "LIST OF SOURCES AND LITERATURE")
    for i, entry in enumerate(bibliography, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Mm(-10)
        p.paragraph_format.left_indent = Mm(10)
        run = p.add_run(f"{i}. {entry['formatted']}")
        _set_run_font(run, Pt(12))


def _add_appendices(doc: Document, appendices: dict[str, str]) -> None:
    # Always renders all 5 appendix slots; empty slots get a placeholder prompt
    _add_chapter_heading(doc, "APPENDICES")
    for i in range(1, 6):
        key = f"appendix_{i}"
        content = appendices.get(key, f"[Appendix {i} — add statistical tables, charts, or documents here]")
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = (i > 1)
        run = p.add_run(f"APPENDIX {i}")
        _set_run_font(run, Pt(12), bold=True)
        _add_body_paragraph(doc, content)


def _add_abstract_placeholder(doc: Document) -> None:
    _add_chapter_heading(doc, "ABSTRACT")
    _add_body_paragraph(
        doc,
        "[Abstract — 0.5–1 page summary of the paper. "
        "Complete this section in Latvian before submission.]"
    )


def _add_originality_declaration(doc: Document) -> None:
    doc.add_page_break()
    _add_chapter_heading(doc, "ORIGINALITY DECLARATION")
    _add_body_paragraph(doc, "Riga, ___ __________ 2026")
    _add_body_paragraph(
        doc,
        "I hereby certify with my signature that the qualification paper submitted "
        "to the Juridical College is an original work and is not plagiarism."
    )
    _add_body_paragraph(doc, "Author: _________________ [name, surname]")
    _add_body_paragraph(doc, "Signature: ______________")


# ── Main assembler ────────────────────────────────────────────────────────────

def build_document(
    chapters: dict[str, str],
    citations: dict,
    appendices: dict[str, str],
) -> Document:
    # Assembly order is fixed by the spec: title → TOC → abbreviations (optional) →
    # chapter headings + sections → bibliography → appendices → abstract → declaration
    doc = Document()
    _apply_page_setup(doc)

    _add_title_page(doc)
    _add_toc_placeholder(doc)

    # Abbreviation list is only inserted when there are 7 or more distinct abbreviations
    abbreviations: list[str] = citations.get("abbreviations", [])
    if len(abbreviations) >= 7:
        _add_abbreviations(doc, abbreviations)

    # Iterate config.SECTIONS (not chapters.keys()) to guarantee chapter heading order
    written_chapter_headings: set[int] = set()
    for section_id in config.SECTIONS:
        if section_id not in chapters:
            continue
        meta = config.SECTION_META.get(section_id, {})
        chapter_num = meta.get("chapter")
        if chapter_num and chapter_num not in written_chapter_headings:
            _add_chapter_heading(doc, config.CHAPTER_HEADINGS[chapter_num])
            written_chapter_headings.add(chapter_num)
        _add_section_content(doc, section_id, chapters[section_id])

    _add_bibliography(doc, citations.get("bibliography", []))
    _add_appendices(doc, appendices)
    _add_abstract_placeholder(doc)
    _add_originality_declaration(doc)

    return doc
