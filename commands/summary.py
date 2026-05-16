from __future__ import annotations
import re
import json
import click
import config
from core.api_client import call
from core.markers import CITE_RE as _CITE_RE
from core.source_manager import load_sources
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from templates.docx_styles import STYLES
from formatters.word_footnotes import FootnoteManager

DRAFTED_SECTIONS: list[str] = [
    "chapter_1_1", "chapter_1_2",
    "chapter_2_1", "chapter_2_2",
    "chapter_3_1", "chapter_3_2",
    "chapter_4_1",
]
NAME_ONLY_SECTIONS: list[str] = ["chapter_4_2", "conclusions"]

_SUMMARISE_SYSTEM = (
    "You are an expert academic editor specialising in AI and human resource management. "
    "Write chapter summaries that read as genuine scholarly prose, not mechanical abstracts. "
    "Each summary must: capture the chapter's central argument and its significance; "
    "vary sentence length and structure; avoid formulaic openers such as "
    "'This chapter examines' or 'This section analyses'; "
    "synthesise topics into a coherent analytical point rather than listing them; "
    "use formal third-person academic English with no em dashes and no first-person pronouns. "
    "Write as many sentences as the chapter's substance requires — typically 2 to 4."
)


def strip_cite_markers(text: str) -> str:
    text = re.sub(r" \{\{cite:[^}]+\}\}(?=\s)", "", text)
    text = re.sub(r"\{\{cite:[^}]+\}\}", "", text)
    return text.strip()


def _load_intro_text() -> str:
    return strip_cite_markers(
        (config.CHAPTERS_DIR / "introduction.md").read_text(encoding="utf-8")
    )


def _make_fn_text_en(src: dict | None, page_ref: str) -> str:
    """Format a footnote with the specific cited page per institution bibliography rules."""
    from formatters.bibliography import (
        format_article, format_article_4_plus,
        format_internet, format_legislation, format_court, EN_DASH,
    )
    page_ref = page_ref.strip()
    m = re.match(r"p\.(\d+)$", page_ref)
    lv_page = f"{m.group(1)}. lpp." if m else page_ref

    if src is None:
        return lv_page or ""

    subtype = src.get("source_subtype", "")
    author = src.get("author", "")
    initial = src.get("initial", "")
    title = src.get("title", "")
    year = src.get("year", 0)
    authors = src.get("authors", [])
    author_count = len(authors) if authors else (1 if author else 0)

    if subtype == "article":
        journal = src.get("journal", "")
        issue = src.get("issue", "")
        if author_count >= 4:
            return format_article_4_plus(title, author, initial, journal, issue, year, lv_page)
        return format_article(author, initial, title, journal, issue, year, lv_page)

    if subtype in ("book", "monograph"):
        city = src.get("city", "")
        publisher = src.get("publisher", "")
        if author_count >= 4:
            return (
                f"{title}/ {author} {initial}. u. c.{EN_DASH} "
                f"{city}: {publisher}, {year}, {lv_page}"
            )
        return f"{author} {initial}. {title}.{EN_DASH} {city}: {publisher}, {year}, {lv_page}"

    if subtype == "legislation":
        return format_legislation(
            title=title,
            adopted_date=src.get("adopted_date", ""),
            publication=src.get("publication", ""),
            pub_ref=src.get("pub_ref", ""),
            pub_date=src.get("pub_date", ""),
        )

    if subtype == "court":
        return format_court(
            case_title=title,
            date=src.get("judgment_date", ""),
            court=src.get("court", ""),
            case_no=src.get("case_no", ""),
            url=src.get("url", ""),
            accessed=src.get("accessed", ""),
        )

    return format_internet(author, initial, title, src.get("url", ""), src.get("accessed", ""))


def _add_paragraph_with_markers(
    doc: Document,
    raw_text: str,
    fm: FootnoteManager,
    sources_map: dict,
    *,
    bold: bool = False,
    font_size=None,
    alignment=None,
    space_before=None,
    first_line_indent=None,
):
    """Add a paragraph, replacing {{cite:}} markers with real Word footnotes."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = STYLES["body"]["line_spacing"]
    if space_before is not None:
        fmt.space_before = space_before
    if first_line_indent is not None:
        fmt.first_line_indent = first_line_indent
    if alignment is not None:
        p.alignment = alignment

    eff_size = font_size or STYLES["body"]["size"]
    pos = 0
    for m in _CITE_RE.finditer(raw_text):
        before = raw_text[pos:m.start()].rstrip(" ")
        if before:
            r = p.add_run(before)
            r.font.name = STYLES["body"]["font"]
            r.font.size = eff_size
            r.bold = bold
        src_id, page_ref = m.group(1), m.group(2) or ""
        fm.add(p._p, src_id, page_ref,
               fn_text=_make_fn_text_en(sources_map.get(src_id), page_ref))
        pos = m.end()

    tail = raw_text[pos:]
    if tail:
        r = p.add_run(tail)
        r.font.name = STYLES["body"]["font"]
        r.font.size = eff_size
        r.bold = bold
    return p


def _load_chapter_text(section: str) -> str:
    return strip_cite_markers(
        (config.CHAPTERS_DIR / f"{section}.md").read_text(encoding="utf-8")
    )


def generate_summaries(sections: list[str]) -> dict[str, str]:
    chapters_block = "\n\n".join(
        f'<chapter id="{s}" title="{config.SECTION_META[s]["title"]}">\n'
        f"{_load_chapter_text(s)}\n</chapter>"
        for s in sections
    )
    user = (
        "For each chapter below, write a summary that captures its central argument and "
        "explains why it matters. Synthesise the content into a unified analytical statement "
        "— do not list topics. Vary sentence openings and lengths. "
        "Return a JSON object mapping each chapter id to its summary string. "
        'Example: {"chapter_1_1": "The integration of AI into HR management has followed '
        'a trajectory shaped as much by institutional readiness as by technological advance."}'
        "\n\n" + chapters_block
    )
    raw = call(system=_SUMMARISE_SYSTEM, user=user, mode="evaluate")
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    try:
        return json.loads(match.group() if match else raw)
    except json.JSONDecodeError as exc:
        click.echo(f"[summary] Failed to parse API response as JSON: {exc}", err=True)
        click.echo(f"[summary] Raw response (first 300 chars): {raw[:300]}", err=True)
        raise


def get_intro_sources() -> list[dict]:
    intro_raw = (config.CHAPTERS_DIR / "introduction.md").read_text(encoding="utf-8")
    ids = list(dict.fromkeys(re.findall(r"\{\{cite:(src_\d+):[^}]+\}\}", intro_raw)))
    src_map = {s["id"]: s for s in load_sources()}
    return [src_map[i] for i in ids if i in src_map]


def format_bib_entry(src: dict) -> str:
    from formatters.bibliography import format_source
    if src.get("source_subtype"):
        return format_source(src)
    parts = []
    if src.get("author"):
        parts.append(src["author"] + ".")
    if src.get("title"):
        parts.append(src["title"] + ".")
    if src.get("publisher"):
        parts.append(src["publisher"] + ",")
    if src.get("year"):
        parts.append(str(src["year"]) + ".")
    if src.get("url"):
        parts.append(src["url"])
    return " ".join(parts)


def _add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    font_size=None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    space_before: Pt | None = None,
    first_line_indent: Mm | None = None,
):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = STYLES["body"]["line_spacing"]
    if space_before is not None:
        fmt.space_before = space_before
    if first_line_indent is not None:
        fmt.first_line_indent = first_line_indent
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = STYLES["body"]["font"]
    run.font.size = font_size or STYLES["body"]["size"]
    run.bold = bold
    return p


def _setup_page_layout(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = STYLES["page"]["width"]
    sec.page_height = STYLES["page"]["height"]
    sec.left_margin = STYLES["page"]["margin_left"]
    sec.right_margin = STYLES["page"]["margin_right"]
    sec.top_margin = STYLES["page"]["margin_top"]
    sec.bottom_margin = STYLES["page"]["margin_bottom"]


def _add_intro_section(
    doc: Document, fm: FootnoteManager, sources_map: dict
) -> None:
    _add_paragraph(
        doc, "INTRODUCTION",
        bold=True, font_size=Pt(14),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=Pt(0),
    )
    raw_intro = (config.CHAPTERS_DIR / "introduction.md").read_text(encoding="utf-8")
    for para in raw_intro.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if re.match(r"^\d+\.", para):
            for line in para.splitlines():
                line = line.strip()
                if line:
                    _add_paragraph_with_markers(
                        doc, line, fm, sources_map,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        first_line_indent=Mm(0),
                    )
        else:
            _add_paragraph_with_markers(
                doc, para, fm, sources_map,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=STYLES["body"]["first_line_indent"],
            )


def _add_chapter_summaries(doc: Document, summaries: dict[str, str]) -> None:
    current_chapter = None
    for section_id in DRAFTED_SECTIONS + NAME_ONLY_SECTIONS:
        meta = config.SECTION_META[section_id]
        chapter_num = meta.get("chapter")
        if chapter_num is not None and chapter_num != current_chapter:
            current_chapter = chapter_num
            chapter_title = f"CHAPTER {chapter_num}. {config.CHAPTER_HEADINGS[chapter_num]}"
            _add_paragraph(
                doc, chapter_title,
                bold=True, font_size=Pt(14),
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(18),
                first_line_indent=Mm(0),
            )
        _add_paragraph(
            doc, meta["heading"],
            bold=True, font_size=Pt(12),
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(6),
            first_line_indent=Mm(0),
        )
        if section_id in summaries:
            _add_paragraph(
                doc, summaries[section_id],
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=STYLES["body"]["first_line_indent"],
            )


def _add_literature_section(doc: Document, sources: list[dict]) -> None:
    _add_paragraph(
        doc, "LITERATURE",
        bold=True, font_size=Pt(14),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=Pt(18),
        first_line_indent=Mm(0),
    )
    for i, src in enumerate(sources, 1):
        _add_paragraph(
            doc, f"{i}. {format_bib_entry(src)}",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Mm(0),
        )


def build_summary_docx(summaries: dict[str, str], sources: list[dict]) -> Document:
    doc = Document()
    _setup_page_layout(doc)
    fm = FootnoteManager(doc)
    sources_map = {s["id"]: s for s in sources if s.get("id")}
    _add_intro_section(doc, fm, sources_map)
    _add_chapter_summaries(doc, summaries)
    _add_literature_section(doc, sources)
    return doc


def run_summary(output_path: str) -> None:
    click.echo("Generating chapter summaries via API...")
    summaries = generate_summaries(DRAFTED_SECTIONS)
    sources = get_intro_sources()
    click.echo(f"Building document ({len(sources)} sources from introduction)...")
    doc = build_summary_docx(summaries, sources)
    doc.save(output_path)
    click.echo(f"Saved: {output_path}")


@click.command()
@click.option("--output", default=None,
              help="Output .docx path (default: papers/<name>/summary_document.docx)")
def summary(output: str | None) -> None:
    """Generate a standalone summary document with introduction and chapter overviews."""
    if output is None:
        output = str(config.PAPER_DIR / "summary_document.docx")
    run_summary(output)
