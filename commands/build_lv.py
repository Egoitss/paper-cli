from __future__ import annotations
import json
import click
import config
from pathlib import Path
from rich.console import Console
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from core.markers import CITE_RE as _CITE_RE
from core.markdown_utils import HEADING_RE as _MD_HEADING_RE, HR_RE as _MD_HR_RE, strip_markdown
from core.text_utils import apply_terminology
from formatters.docx_builder import _apply_page_setup, _set_run_font
from formatters.bibliography import sort_bibliography
from formatters.word_footnotes import FootnoteManager

console = Console()

SECTION_HEADINGS_LV = {
    "introduction": "IEVADS",
    "part_1": "1. TEORĒTISKĀ DAĻA",
    "part_2": "2. PRAKTISKĀ DAĻA",
    "conclusions": "SECINĀJUMI",
}

LV = {
    "institution": "JURIDISKĀ KOLEDŽA",
    "doc_type": "ZIŅOJUMS",
    "author_label": "Autors: _________________________ [vārds, uzvārds, grupas kods]",
    "supervisor_label": "Zinātniskais vadītājs: _______________ [grāds, vārds, uzvārds]",
    "city_year": "RĪGA 2026",
    "toc_title": "SATURS",
    "bibliography_title": "AVOTU SARAKSTS",
    "accessed": "Skatīts",
}

TERMINOLOGY = {"CJEU": "EST", "EDPB": "EDAK"}


# ── Document structure helpers ────────────────────────────────────────────────

def _add_body_paragraph(doc: Document, text: str, fm: FootnoteManager) -> None:
    """Add a body paragraph, inserting Word footnote references for {{cite:...}} markers."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(15)
    p.paragraph_format.line_spacing = 1.5

    last = 0
    for m in _CITE_RE.finditer(text):
        before = strip_markdown(text[last:m.start()])
        if before:
            _set_run_font(p.add_run(before), Pt(12))
        fm.add(p._p, m.group(1), m.group(2) or "")
        last = m.end()

    remaining = strip_markdown(text[last:])
    if remaining:
        _set_run_font(p.add_run(remaining), Pt(12))


def _add_chapter_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_toc(doc: Document) -> None:
    _add_chapter_heading(doc, LV["toc_title"])
    settings_el = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings_el.append(upd)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-1" \\h \\z \\u '
    run._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _add_title_page(doc: Document) -> None:
    title_lv = config.PAPER_TITLE_LV or config.PAPER_TITLE
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(LV["institution"]), Pt(16))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(48)
    _set_run_font(p2.add_run(title_lv.upper()), Pt(24), bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    _set_run_font(p3.add_run(LV["doc_type"]), Pt(16))
    for _ in range(4):
        doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(p4.add_run(LV["author_label"]), Pt(16))
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(p5.add_run(LV["supervisor_label"]), Pt(16))
    for _ in range(4):
        doc.add_paragraph()
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p6.add_run(LV["city_year"]), Pt(16))


def _add_section(doc: Document, section_id: str, text: str, fm: FootnoteManager) -> None:
    heading = SECTION_HEADINGS_LV.get(
        section_id,
        config.SECTION_META.get(section_id, {}).get("heading", section_id.upper()),
    )
    _add_chapter_heading(doc, heading)
    for block in text.split("\n\n"):
        stripped = block.strip()
        if not stripped or _MD_HEADING_RE.match(stripped) or _MD_HR_RE.match(stripped):
            continue
        _add_body_paragraph(doc, stripped, fm)


def _format_source_lv(src: dict) -> str:
    author = src.get("author", "")
    initial = src.get("initial", (author or "")[:1])
    title = src.get("title", "")
    url = src.get("url", "")
    accessed = src.get("accessed", "")
    year = src.get("year", "")
    publisher = src.get("publisher", "")
    sep = " – "
    author_part = (f"{author} {initial}. ".strip(". ") + ". ") if author else ""
    parts = [f"{author_part}{title} ({year})", publisher]
    if url:
        parts.append(url)
    if accessed:
        parts.append(f"({LV['accessed']}: {accessed})")
    return sep.join(p for p in parts if p) + "."


def _add_bibliography(doc: Document, sources: list[dict]) -> None:
    _add_chapter_heading(doc, LV["bibliography_title"])
    for i, src in enumerate(sort_bibliography(sources), 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Mm(-10)
        p.paragraph_format.left_indent = Mm(10)
        p.paragraph_format.line_spacing = 1.5
        _set_run_font(p.add_run(f"{i}. {_format_source_lv(src)}"), Pt(12))


# ── CLI command ───────────────────────────────────────────────────────────────

@click.command("build-lv")
@click.option("--output", default=None,
              help="Output path (default: papers/<name>/submission_draft_lv.docx)")
def build_lv(output: str) -> None:
    """Assemble Latvian _lv.md chapters into a submission-ready .docx file."""
    if output is None:
        output = str(config.PAPER_DIR / "submission_draft_lv.docx")

    sources: list[dict] = json.loads(config.SOURCES_FILE.read_text(encoding="utf-8"))

    raw_sections: dict[str, str] = {}
    for section_id in config.SECTIONS:
        lv_path = config.CHAPTERS_DIR / f"{section_id}_lv.md"
        if lv_path.exists():
            raw_sections[section_id] = apply_terminology(
                lv_path.read_text(encoding="utf-8"), TERMINOLOGY
            )
        else:
            console.print(f"[yellow]Skipping (no _lv.md):[/yellow] {section_id}")

    if not raw_sections:
        console.print("[red]No _lv.md files found.[/red]")
        raise SystemExit(1)

    doc = Document()
    _apply_page_setup(doc)

    fm = FootnoteManager(doc)

    _add_title_page(doc)
    _add_toc(doc)

    for section_id in config.SECTIONS:
        if section_id in raw_sections:
            _add_section(doc, section_id, raw_sections[section_id], fm)

    _add_bibliography(doc, sources)

    out_path = Path(output)
    doc.save(str(out_path))
    console.print(f"[green]Saved:[/green] {out_path}")
    console.print(f"[dim]{fm._next_id - 1} footnotes inserted. "
                  "Word auto-updates TOC on first open.[/dim]")
