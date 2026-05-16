import json
import pytest
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentClass
from formatters.docx_builder import build_document


@pytest.fixture
def minimal_input():
    chapters = {"introduction": "This is the introduction text. No citations here."}
    citations = {"footnotes": [], "bibliography": [], "abbreviations": []}
    appendices = {}
    return chapters, citations, appendices


def test_build_document_returns_document(minimal_input):
    chapters, citations, appendices = minimal_input
    doc = build_document(chapters, citations, appendices)
    assert isinstance(doc, DocumentClass)


def test_build_document_contains_introduction_text(minimal_input):
    chapters, citations, appendices = minimal_input
    doc = build_document(chapters, citations, appendices)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "introduction" in full_text.lower() or "This is the introduction" in full_text


def test_build_document_includes_abbreviations_when_7_or_more(minimal_input):
    chapters, citations, appendices = minimal_input
    citations["abbreviations"] = ["AI", "ML", "NLP", "HRM", "RPA", "ATS", "GDPR"]
    doc = build_document(chapters, citations, appendices)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ABBREVIATIONS" in full_text.upper()


def test_build_document_omits_abbreviations_when_fewer_than_7(minimal_input):
    chapters, citations, appendices = minimal_input
    citations["abbreviations"] = ["AI", "ML"]
    doc = build_document(chapters, citations, appendices)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "LIST OF ABBREVIATIONS" not in full_text.upper()


def test_build_document_margins_correct(minimal_input):
    chapters, citations, appendices = minimal_input
    from docx.shared import Mm
    doc = build_document(chapters, citations, appendices)
    section = doc.sections[0]
    # Tolerance of 914 EMU (~0.025mm) accounts for twips rounding in OOXML storage
    assert abs(section.left_margin - Mm(30)) < 914
    assert abs(section.right_margin - Mm(15)) < 914


def test_build_document_page_numbers_configured(minimal_input):
    chapters, citations, appendices = minimal_input
    doc = build_document(chapters, citations, appendices)
    section = doc.sections[0]
    # Title page counted but "1" not printed
    assert section.different_first_page_header_footer is True
    # PAGE field injected into header
    hdr_xml = section.header._element.xml
    assert "PAGE" in hdr_xml
